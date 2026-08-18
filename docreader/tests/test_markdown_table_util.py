import io
import unittest

from markitdown import MarkItDown
from docx import Document

from docreader.parser.markdown_parser import MarkdownTableUtil


class TestMarkdownTableUtil(unittest.TestCase):
    def test_preserves_empty_cells(self):
        """Interior empty cells must not be dropped during formatting."""
        raw = "| a |  | c |\n| --- | --- | --- |\n| 1 | 2 | 3 |"
        formatted = MarkdownTableUtil().format_table(raw)
        self.assertIn("| a |  | c |", formatted)
        self.assertEqual(formatted.count("|"), raw.count("|"))

    def test_format_nonempty_table(self):
        raw = "|Name|Age|\n|---|---|\n|John|30|"
        formatted = MarkdownTableUtil().format_table(raw)
        self.assertIn("| Name | Age |", formatted)
        self.assertIn("| --- | --- |", formatted)
        self.assertIn("| John | 30 |", formatted)

    def test_normalize_markitdown_en_tables(self):
        docx = io.BytesIO()
        document = Document()
        tables = [
            [
                ["Name", "Game", "Fame", "Blame"],
                ["Lebron James", "Basketball", "", ""],
            ],
            [["Sinple", "Table"], ["Without", "Header"]],
            [
                ["Simple Multiparagraph", "Table Full"],
                ["Of Paragraphs", "In each Cell."],
            ],
        ]
        for rows in tables:
            table = document.add_table(rows=len(rows), cols=len(rows[0]))
            for row, values in zip(table.rows, rows):
                for cell, value in zip(row.cells, values):
                    cell.text = value
            document.add_paragraph("")
        document.save(docx)

        raw = MarkItDown().convert(
            io.BytesIO(docx.getvalue()), file_extension=".docx"
        ).text_content
        normalized = MarkdownTableUtil().format_table(raw)

        self.assertNotIn("|  |  |  |  |", normalized)
        self.assertIn("| Name | Game | Fame | Blame |", normalized)
        idx_name = normalized.index("| Name | Game | Fame | Blame |")
        idx_sep = normalized.index("| --- | --- | --- | --- |", idx_name)
        self.assertLess(idx_name, idx_sep)
        self.assertIn("| Lebron James | Basketball |", normalized)

        # Headerless 2-row tables: delimiter inserted so GFM renderers show a table
        self.assertIn(
            "| Sinple | Table |\n| --- | --- |\n| Without | Header |", normalized
        )
        self.assertIn(
            "| Simple Multiparagraph | Table Full |\n| --- | --- |\n"
            "| Of Paragraphs | In each Cell. |",
            normalized,
        )


if __name__ == "__main__":
    unittest.main()
